from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os
import re
import spacy
import textstat
import unicodedata
from datetime import datetime
import json

# Load spaCy model once, with minimal components for consistency
nlp = spacy.load("en_core_web_sm", disable=["ner", "lemmatizer"])

def extract_text(file):
    """Extract text from PDF or DOCX file with normalization."""
    try:
        text = ""
        if file.filename.endswith(".pdf"):
            reader = PdfReader(file.file)
            text = "".join(page.extract_text() or "" for page in reader.pages)
        elif file.filename.endswith((".doc", ".docx")):
            file.file.seek(0)
            doc = Document(file.file)
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        else:
            raise ValueError("Unsupported file format")

        # Normalize text to remove encoding issues and extra whitespace
        text = unicodedata.normalize("NFKD", text)
        text = re.sub(r'\s+', ' ', text.strip())
        if not text:
            raise ValueError("No text extracted from file")
        print(f"Extracted text (first 200 chars): {text[:200]}")  # Debug
        return text
    except Exception as e:
        raise Exception(f"Failed to parse resume: {str(e)}")

def extract_resume_sections(text):
    """Extract structured information from resume text."""
    sections = {
        'contact': {},
        'summary': '',
        'experience': [],
        'education': [],
        'skills': [],
        'projects': [],
        'achievements': []
    }
    
    # Extract contact information
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
    
    if emails:
        sections['contact']['email'] = emails[0]
    if phones:
        sections['contact']['phone'] = ''.join(phones[0])
    if linkedin:
        sections['contact']['linkedin'] = linkedin[0]
    
    # Extract summary/objective
    summary_patterns = [
        r'(?:summary|objective|profile|about)[\s:]*([^.]+(?:\.[^.]*)*)',
        r'(?:professional\s+summary|career\s+objective)[\s:]*([^.]+(?:\.[^.]*)*)'
    ]
    
    for pattern in summary_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            sections['summary'] = match.group(1).strip()
            break
    
    # Extract experience
    exp_pattern = r'(?:experience|work\s+history|employment)[\s:]*([\s\S]*?)(?=education|skills|projects|$|education|skills|projects)'
    exp_match = re.search(exp_pattern, text, re.IGNORECASE | re.DOTALL)
    if exp_match:
        exp_text = exp_match.group(1)
        # Simple extraction of job entries
        job_entries = re.split(r'\n\s*\n', exp_text)
        for entry in job_entries:
            if entry.strip() and len(entry.strip()) > 20:
                sections['experience'].append(entry.strip())
    
    # Extract education
    edu_pattern = r'(?:education|academic|qualification)[\s:]*([\s\S]*?)(?=experience|skills|projects|$|experience|skills|projects)'
    edu_match = re.search(edu_pattern, text, re.IGNORECASE | re.DOTALL)
    if edu_match:
        edu_text = edu_match.group(1)
        edu_entries = re.split(r'\n\s*\n', edu_text)
        for entry in edu_entries:
            if entry.strip() and len(entry.strip()) > 10:
                sections['education'].append(entry.strip())
    
    # Extract skills
    skills_patterns = [
        r'(?:skills|technical\s+skills|competencies)[\s:]*([\s\S]*?)(?=experience|education|projects|$|experience|education|projects)',
        r'(?:programming\s+languages|technologies)[\s:]*([\s\S]*?)(?=experience|education|projects|$|experience|education|projects)'
    ]
    
    for pattern in skills_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            skills_text = match.group(1)
            # Extract skills separated by commas, semicolons, or newlines
            skills = re.split(r'[,;\n]', skills_text)
            for skill in skills:
                skill = skill.strip()
                if skill and len(skill) > 1:
                    sections['skills'].append(skill)
            break
    
    return sections

def compute_resume_score(text):
    """Compute comprehensive professionalism score (0-100) with detailed analysis."""
    try:
        score_dict = {
            "structure": 0,
            "grammar": 0,
            "readability": 0,
            "keywords": 0,
            "length": 0,
            "contact": 0,
            "achievements": 0,
            "formatting": 0,
            "action_verbs": 0,
            "quantification": 0
        }

        # Extract structured sections
        sections = extract_resume_sections(text)
        
        # Structure: Count sections and bullets (25 points)
        standard_sections = ['summary', 'experience', 'education', 'skills']
        found_sections = sum(1 for section in standard_sections if sections[section])
        bullets = len(re.findall(r'^\s*[-•*]\s', text, re.MULTILINE))
        score_dict["structure"] = min((found_sections * 5 + bullets * 1), 25)

        # Grammar assessment (20 points)
        doc = nlp(text[:10000])
        grammar_errors = 0
        for sent in doc.sents:
            stripped = sent.text.strip()
            if stripped and not stripped[0].isupper():
                grammar_errors += 1
            if len(stripped) > 30 and not stripped.endswith(('.', '!', '?')):
                grammar_errors += 1
        score_dict["grammar"] = max(20 - min(grammar_errors, 20), 0)

        # Readability (15 points)
        readability = textstat.flesch_reading_ease(text)
        score_dict["readability"] = min(readability / 6.67, 15)

        # Keywords and skills (15 points)
        tech_keywords = [
            "python", "javascript", "react", "sql", "management", "node", "java", 
            "docker", "kubernetes", "aws", "azure", "gcp", "machine learning", 
            "data analysis", "project management", "agile", "scrum", "git", 
            "rest api", "microservices", "devops", "ci/cd", "tensorflow", "pytorch"
        ]
        keywords_found = sum(1 for word in tech_keywords if word.lower() in text.lower())
        score_dict["keywords"] = min(keywords_found * 2, 15)

        # Length optimization (10 points)
        word_count = len(text.split())
        if 300 <= word_count <= 800:
            score_dict["length"] = 10
        else:
            score_dict["length"] = max(10 - abs(word_count - 550) // 50, 0)

        # Contact information (10 points)
        contact_score = 0
        if sections['contact'].get('email'):
            contact_score += 4
        if sections['contact'].get('phone'):
            contact_score += 3
        if sections['contact'].get('linkedin'):
            contact_score += 3
        score_dict["contact"] = contact_score

        # Achievements and impact (5 points)
        achievement_words = ['achieved', 'increased', 'improved', 'reduced', 'developed', 'created', 'managed', 'led', 'implemented']
        achievement_count = sum(1 for word in achievement_words if word.lower() in text.lower())
        score_dict["achievements"] = min(achievement_count, 5)

        # Formatting and presentation (10 points)
        formatting_score = 0
        # Check for consistent formatting
        if re.search(r'\b[A-Z][a-z]+\s+[A-Z][a-z]+', text):  # Proper name formatting
            formatting_score += 2
        if re.search(r'\d{4}', text):  # Has years/dates
            formatting_score += 2
        if len(re.findall(r'[A-Z][a-z]+\s+[A-Z][a-z]+', text)) > 2:  # Multiple proper nouns
            formatting_score += 2
        if re.search(r'@', text):  # Has email
            formatting_score += 2
        if re.search(r'\(\d{3}\)\s*\d{3}-\d{4}|\d{3}-\d{3}-\d{4}', text):  # Has phone
            formatting_score += 2
        score_dict["formatting"] = min(formatting_score, 10)

        # Action verbs usage (5 points)
        action_verbs = [
            'achieved', 'accomplished', 'administered', 'analyzed', 'assisted', 'built', 'collaborated',
            'created', 'delivered', 'designed', 'developed', 'executed', 'facilitated', 'generated',
            'implemented', 'improved', 'increased', 'initiated', 'launched', 'led', 'managed',
            'optimized', 'organized', 'performed', 'planned', 'produced', 'reduced', 'resolved',
            'streamlined', 'supervised', 'transformed', 'utilized'
        ]
        action_verb_count = sum(1 for verb in action_verbs if verb.lower() in text.lower())
        score_dict["action_verbs"] = min(action_verb_count, 5)

        # Quantification and metrics (5 points)
        # Look for numbers, percentages, dollar amounts, time periods
        numbers = re.findall(r'\b\d+(?:\.\d+)?%?\b', text)
        quantified_phrases = re.findall(r'\b(?:increased|decreased|reduced|improved|saved|generated|managed|led|supervised|trained|developed|created|built|delivered|achieved|accomplished|grew|expanded|optimized|streamlined|enhanced|boosted|raised|lowered|cut|eliminated|reduced|minimized|maximized|doubled|tripled|quadrupled|halved|by\s+\d+%?|\d+x|\d+\s+times|\d+\s+fold)\b', text, re.IGNORECASE)
        quantification_score = min(len(numbers) + len(quantified_phrases), 5)
        score_dict["quantification"] = quantification_score

        total_score = sum(score_dict.values())
        total_score = int(min(max(total_score, 0), 100))

        # Build comprehensive issue list
        issues = []
        
        # Structure issues
        if found_sections < 3:
            issues.append("Add standard sections: Summary, Experience, Education, Skills.")
        if bullets < 5:
            issues.append("Use bullet points to highlight achievements and responsibilities.")
        if not sections['summary']:
            issues.append("Add a professional summary or objective statement.")
        
        # Grammar issues
        if grammar_errors > 0:
            issues.append(f"Fix capitalization and sentence punctuation in {grammar_errors} places.")
        
        # Readability issues
        if readability < 50:
            issues.append("Simplify sentences to improve readability (aim for 60+ Flesch score).")
        
        # Keyword issues
        if keywords_found < 3:
            issues.append("Include more role-relevant keywords and technical skills.")
        
        # Length issues
        if word_count < 300:
            issues.append("Resume is too short; add details on projects, achievements, and impact.")
        elif word_count > 800:
            issues.append("Resume is too long; trim to most relevant achievements and experiences.")
        
        # Contact issues
        if not sections['contact'].get('email'):
            issues.append("Add a professional email address.")
        if not sections['contact'].get('phone'):
            issues.append("Include a contact phone number.")
        if not sections['contact'].get('linkedin'):
            issues.append("Add your LinkedIn profile URL.")
        
        # Achievement issues
        if achievement_count < 2:
            issues.append("Include more quantifiable achievements and impact statements.")
        
        # Experience issues
        if len(sections['experience']) < 2:
            issues.append("Add more detailed work experience with specific achievements.")
        
        # Skills issues
        if len(sections['skills']) < 5:
            issues.append("Expand your skills section with relevant technical and soft skills.")

        # Formatting issues
        if score_dict["formatting"] < 6:
            issues.append("Improve formatting consistency and professional presentation.")
        
        # Action verb issues
        if score_dict["action_verbs"] < 3:
            issues.append("Use more strong action verbs to describe your accomplishments.")
        
        # Quantification issues
        if score_dict["quantification"] < 2:
            issues.append("Add more numbers, percentages, and metrics to quantify your achievements.")

        return total_score, issues, sections, score_dict
    except Exception as e:
        raise Exception(f"Score computation failed: {str(e)}")


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink style
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)

def create_professional_resume(sections, output_dir, original_filename):
    """Create a professionally formatted resume with improved structure and styling."""
    try:
        doc = Document()
        
        # Set document margins
        sections_obj = doc.sections[0]
        sections_obj.top_margin = Inches(0.5)
        sections_obj.bottom_margin = Inches(0.5)
        sections_obj.left_margin = Inches(0.7)
        sections_obj.right_margin = Inches(0.7)
        
        # Header Section
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Name (extract from filename or use default)
        name = os.path.splitext(original_filename)[0].replace('_', ' ').replace('-', ' ').title()
        name_run = header.add_run(name)
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        
        # Contact Information
        if sections.get('contact'):
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            contact_info = []
            if sections['contact'].get('email'):
                contact_info.append(sections['contact']['email'])
            if sections['contact'].get('phone'):
                contact_info.append(sections['contact']['phone'])
            if sections['contact'].get('linkedin'):
                contact_info.append(sections['contact']['linkedin'])
            
            if contact_info:
                contact_text = ' | '.join(contact_info)
                contact_run = contact_para.add_run(contact_text)
                contact_run.font.size = Pt(10)
        
        # Add line separator
        doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Professional Summary
        if sections.get('summary'):
            add_section_header(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(sections['summary'])
            summary_run.font.size = Pt(10)
            summary_para.paragraph_format.space_after = Pt(6)
        
        # Experience Section
        if sections.get('experience') and len(sections['experience']) > 0:
            add_section_header(doc, "PROFESSIONAL EXPERIENCE")
            for exp in sections['experience'][:3]:  # Limit to top 3 experiences
                exp_para = doc.add_paragraph()
                exp_run = exp_para.add_run(str(exp))
                exp_run.font.size = Pt(10)
                exp_para.paragraph_format.space_after = Pt(3)
                exp_para.paragraph_format.left_indent = Inches(0.2)
        
        # Education Section
        if sections.get('education') and len(sections['education']) > 0:
            add_section_header(doc, "EDUCATION")
            for edu in sections['education'][:2]:  # Limit to top 2 education entries
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(str(edu))
                edu_run.font.size = Pt(10)
                edu_para.paragraph_format.space_after = Pt(3)
                edu_para.paragraph_format.left_indent = Inches(0.2)
        
        # Skills Section
        if sections.get('skills') and len(sections['skills']) > 0:
            add_section_header(doc, "TECHNICAL SKILLS")
            skills_para = doc.add_paragraph()
            skills_text = ' • '.join(sections['skills'][:15])  # Limit to top 15 skills
            skills_run = skills_para.add_run(skills_text)
            skills_run.font.size = Pt(10)
            skills_para.paragraph_format.space_after = Pt(6)
            skills_para.paragraph_format.left_indent = Inches(0.2)
        
        # Add footer with generation date
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')} by SGP Resume Analyzer")
        footer_run.font.size = Pt(8)
        footer_run.italic = True
        
        # Save the document
        os.makedirs(output_dir, exist_ok=True)
        safe_base = os.path.splitext(os.path.basename(original_filename))[0]
        out_name = f"professional-{safe_base}.docx"
        out_path = os.path.join(output_dir, out_name)
        doc.save(out_path)
        print(f"Professional resume saved to: {out_path}")
        return out_path
        
    except Exception as e:
        print(f"Error creating professional resume: {str(e)}")
        raise Exception(f"Failed to create professional resume: {str(e)}")

def add_section_header(doc, title):
    """Add a formatted section header to the document."""
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(title)
    header_run.font.size = Pt(12)
    header_run.font.bold = True
    header_para.paragraph_format.space_before = Pt(8)
    header_para.paragraph_format.space_after = Pt(3)

def generate_corrected_docx_preserving_layout(original_file, output_dir):
    """Light-touch corrections that keep layout: normalize spaces in runs, ensure minimal spacing, keep fonts when set."""
    try:
        original_file.file.seek(0)
        doc = Document(original_file.file)
        for p in doc.paragraphs:
            # normalize spacing inside each run, preserve bold/italic and font
            for run in p.runs:
                if run.text:
                    run.text = re.sub(r"\s+", " ", run.text)
                    run.text = re.sub(r"\.([A-Za-z])", r". \1", run.text)
            if p.paragraph_format.space_after is None:
                p.paragraph_format.space_after = Pt(2)

        os.makedirs(output_dir, exist_ok=True)
        safe_base = os.path.splitext(os.path.basename(original_file.filename))[0]
        out_name = f"updated-{safe_base}.docx"
        out_path = os.path.join(output_dir, out_name)
        doc.save(out_path)
        return out_path
    except Exception as e:
        raise Exception(f"Failed to generate corrected DOCX: {str(e)}")

def generate_enhanced_resume(original_file, output_dir, sections):
    """Generate both corrected and professional versions of the resume."""
    try:
        # Generate corrected version (preserving original layout)
        corrected_path = generate_corrected_docx_preserving_layout(original_file, output_dir)
        
        # Generate professional version
        professional_path = create_professional_resume(sections, output_dir, original_file.filename)
        
        return corrected_path, professional_path
    except Exception as e:
        raise Exception(f"Failed to generate enhanced resume: {str(e)}")